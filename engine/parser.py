"""
Resume parsing module for the Internship Recommendation Engine.

Extracts raw text from uploaded resume files (.pdf, .docx, .txt) and returns
a cleaned, normalised string ready for downstream NLP processing. Scanned /
image-based PDF pages are handled transparently via OCR fallback.

Supported formats:
    - **.pdf**  — text extraction with ``pdfplumber``; per-page OCR fallback
      using ``pdf2image`` + ``pytesseract`` when a page yields no text.
    - **.docx** — paragraph and table-cell extraction with ``python-docx``.
    - **.txt**  — plain-text read with automatic encoding detection.

Raises:
    ResumeParseError: On file-not-found, unsupported format, or empty result.
"""

from __future__ import annotations

import logging
import os
import re
import sys
from pathlib import Path
from typing import Final

logger = logging.getLogger(__name__)

# ── File extension constants ─────────────────────────────────────────────────

_PDF_EXT: Final[str] = ".pdf"
_DOCX_EXT: Final[str] = ".docx"
_TXT_EXT: Final[str] = ".txt"
_SUPPORTED_EXTENSIONS: Final[frozenset[str]] = frozenset({_PDF_EXT, _DOCX_EXT, _TXT_EXT})

# Regex: Unicode control chars (C0/C1) except \n, \r, \t
_CONTROL_CHAR_RE: Final[re.Pattern[str]] = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]"
)


# ── Custom exception ─────────────────────────────────────────────────────────


class ResumeParseError(Exception):
    """Raised when resume text extraction fails irrecoverably.

    Common causes:
        - The file does not exist on disk.
        - The file extension is not one of .pdf, .docx, or .txt.
        - All extraction strategies (including OCR) returned empty text.
    """


# ── Internal helpers ─────────────────────────────────────────────────────────


def _clean_text(raw: str) -> str:
    """Strip control characters and collapse excessive whitespace.

    Args:
        raw: The raw extracted text, potentially containing artefacts from
            PDF extraction or OCR.

    Returns:
        Cleaned text with control characters removed, runs of whitespace
        collapsed to single spaces, and leading/trailing whitespace stripped.
        Paragraph breaks (double newlines) are preserved.
    """
    # Remove control characters (keep \n, \r, \t)
    text = _CONTROL_CHAR_RE.sub("", raw)

    # Normalise line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Collapse runs of blank lines into a single paragraph break
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Collapse horizontal whitespace (spaces, tabs) within lines
    text = re.sub(r"[ \t]+", " ", text)

    # Strip leading/trailing whitespace on each line
    text = "\n".join(line.strip() for line in text.splitlines())

    return text.strip()


def _extract_pdf(file_path: Path) -> str:
    """Extract text from a PDF, falling back to OCR for scanned pages.

    Uses ``pdfplumber`` for native text extraction. If a page returns empty
    or whitespace-only text, the page is rendered to an image via
    ``pdf2image`` and processed with ``pytesseract``.

    Args:
        file_path: Absolute path to the .pdf file.

    Returns:
        Concatenated text from all pages.

    Raises:
        ResumeParseError: If pdfplumber fails to open the file.
    """
    import pdfplumber  # lazy import to keep startup fast

    pages_text: list[str] = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""

                if page_text.strip():
                    logger.debug("Page %d: extracted %d chars via pdfplumber.", page_num, len(page_text))
                    pages_text.append(page_text)
                else:
                    # Scanned / image page → OCR fallback
                    logger.info("Page %d: no text found, attempting OCR fallback.", page_num)
                    ocr_text = _ocr_pdf_page(file_path, page_num)
                    if ocr_text.strip():
                        pages_text.append(ocr_text)
                    else:
                        logger.warning("Page %d: OCR also returned empty text.", page_num)
    except Exception as exc:
        raise ResumeParseError(
            f"Failed to read PDF '{file_path.name}': {exc}"
        ) from exc

    return "\n\n".join(pages_text)


def _ocr_pdf_page(file_path: Path, page_number: int) -> str:
    """Run OCR on a single PDF page using pdf2image + pytesseract.

    Args:
        file_path: Absolute path to the .pdf file.
        page_number: 1-indexed page number to OCR.

    Returns:
        OCR-extracted text for the requested page.
    """
    from pdf2image import convert_from_path  # lazy import
    import pytesseract  # lazy import

    try:
        images = convert_from_path(
            str(file_path),
            first_page=page_number,
            last_page=page_number,
            dpi=300,
        )
    except Exception as exc:
        logger.error("pdf2image failed on page %d of '%s': %s", page_number, file_path.name, exc)
        return ""

    if not images:
        return ""

    text: str = pytesseract.image_to_string(images[0], lang="eng")
    logger.debug("OCR page %d: extracted %d chars.", page_number, len(text))
    return text


def _extract_docx(file_path: Path) -> str:
    """Extract text from a .docx file (paragraphs and table cells).

    Args:
        file_path: Absolute path to the .docx file.

    Returns:
        Concatenated text from all paragraphs followed by all table cells.

    Raises:
        ResumeParseError: If python-docx fails to open the file.
    """
    from docx import Document  # lazy import

    try:
        doc = Document(str(file_path))
    except Exception as exc:
        raise ResumeParseError(
            f"Failed to read DOCX '{file_path.name}': {exc}"
        ) from exc

    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Table cells (may contain additional structured info like education rows)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return "\n".join(parts)


def _extract_txt(file_path: Path) -> str:
    """Read a plain-text file with automatic encoding fallback.

    Tries UTF-8 first, then falls back to latin-1 which never raises a
    decoding error.

    Args:
        file_path: Absolute path to the .txt file.

    Returns:
        The full file content as a string.

    Raises:
        ResumeParseError: If the file cannot be read.
    """
    for encoding in ("utf-8", "latin-1"):
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            raise ResumeParseError(
                f"Failed to read TXT '{file_path.name}': {exc}"
            ) from exc

    # Should never reach here because latin-1 accepts all byte values
    raise ResumeParseError(f"Unable to decode '{file_path.name}' with any supported encoding.")


# ── Public API ───────────────────────────────────────────────────────────────


def parse_resume(file_path: str) -> str:
    """Parse a resume file and return cleaned plain text.

    This is the main entry point for the parsing module. It resolves the file
    path, selects the appropriate extraction strategy based on extension,
    cleans the result, and raises :class:`ResumeParseError` if no usable text
    can be produced.

    Args:
        file_path: Path (absolute or relative) to the resume file.

    Returns:
        Cleaned, normalised plain-text content of the resume.

    Raises:
        ResumeParseError: If the file does not exist, the format is
            unsupported, or the extracted text is empty after all attempts.

    Examples:
        >>> text = parse_resume("resumes/jane_doe.pdf")
        >>> len(text) > 0
        True
    """
    path = Path(file_path).resolve()

    # ── Validate existence ────────────────────────────────────────────────
    if not path.exists():
        raise ResumeParseError(f"File not found: {path}")

    if not path.is_file():
        raise ResumeParseError(f"Path is not a file: {path}")

    # ── Validate extension ────────────────────────────────────────────────
    ext = path.suffix.lower()
    if ext not in _SUPPORTED_EXTENSIONS:
        raise ResumeParseError(
            f"Unsupported file format '{ext}'. "
            f"Supported formats: {', '.join(sorted(_SUPPORTED_EXTENSIONS))}"
        )

    logger.info("Parsing resume '%s' (format: %s).", path.name, ext)

    # ── Dispatch to handler ───────────────────────────────────────────────
    if ext == _PDF_EXT:
        raw_text = _extract_pdf(path)
    elif ext == _DOCX_EXT:
        raw_text = _extract_docx(path)
    else:
        raw_text = _extract_txt(path)

    # ── Post-process ──────────────────────────────────────────────────────
    cleaned = _clean_text(raw_text)

    if not cleaned:
        raise ResumeParseError(
            f"Extracted text from '{path.name}' is empty after cleaning. "
            "The file may be a scanned image without recognisable text, or "
            "the document may genuinely contain no textual content."
        )

    logger.info("Successfully parsed '%s': %d characters extracted.", path.name, len(cleaned))
    return cleaned


# ── CLI entry point ──────────────────────────────────────────────────────────


if __name__ == "__main__":
    logging.basicConfig(
        level=logging.DEBUG,
        format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
    )

    if len(sys.argv) < 2:
        print("Usage: python -m engine.parser <resume_file_path>")
        sys.exit(1)

    target = sys.argv[1]

    try:
        result = parse_resume(target)
        print(f"\n{'=' * 72}")
        print(f"  Parsed resume: {os.path.basename(target)}")
        print(f"  Characters   : {len(result)}")
        print(f"{'=' * 72}\n")
        print(result[:3000])  # preview first 3 000 chars
        if len(result) > 3000:
            print(f"\n... [{len(result) - 3000} more characters truncated]")
    except ResumeParseError as e:
        print(f"\n❌ ResumeParseError: {e}", file=sys.stderr)
        sys.exit(1)
